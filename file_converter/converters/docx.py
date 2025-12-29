"""
file_converter/converters/docx.py
DOCX 文件转换器 - 使用 python-docx 提取文档内容，支持嵌入图片 OCR
"""

import io
import logging
from typing import BinaryIO, List, Optional

from docx import Document
from docx.shared import Pt

from file_converter.converters.base import BaseConverter, ConversionResult

logger = logging.getLogger(__name__)


class DocxConverter(BaseConverter):
    """
    DOCX 转 Markdown 转换器
    
    功能：
    - 保留 Heading 1-6 层级转换为 # 标记
    - 保留段落、列表、表格结构
    - 加粗、斜体转换为 Markdown 语法
    - 图片标注为 [图片略]
    - 支持嵌入图片 OCR 识别
    """
    
    supported_extensions = ["docx", "doc"]
    
    # Word 标题样式到 Markdown 层级的映射
    HEADING_STYLES = {
        'Heading 1': 1, 'Heading1': 1, 'heading 1': 1,
        'Heading 2': 2, 'Heading2': 2, 'heading 2': 2,
        'Heading 3': 3, 'Heading3': 3, 'heading 3': 3,
        'Heading 4': 4, 'Heading4': 4, 'heading 4': 4,
        'Heading 5': 5, 'Heading5': 5, 'heading 5': 5,
        'Heading 6': 6, 'Heading6': 6, 'heading 6': 6,
        'Title': 1, 'Subtitle': 2,
    }
    
    def convert(
        self, 
        file: BinaryIO, 
        filename: str,
        enable_ocr: bool = False,
        ocr_lang: Optional[str] = None
    ) -> ConversionResult:
        """将 DOCX 转换为 Markdown"""
        warnings: List[str] = []
        content_parts: List[str] = []
        ocr_count = 0
        
        try:
            doc = Document(file)
            logger.info(f"Processing DOCX: {filename}, enable_ocr={enable_ocr}")
            
            # 处理段落
            for para in doc.paragraphs:
                para_md = self._process_paragraph(para)
                if para_md:
                    content_parts.append(para_md)
                
                # 检查段落中的图片
                if enable_ocr:
                    images = para._element.xpath('.//w:drawing')
                    if images:
                        ocr_text = self._ocr_paragraph_images(para, warnings, ocr_lang)
                        if ocr_text:
                            content_parts.append(f"> [图片文字] {ocr_text}")
                            ocr_count += 1
                        else:
                            content_parts.append("[图片略]")
                else:
                    if para._element.xpath('.//w:drawing'):
                        content_parts.append("[图片略]")
            
            # 处理表格
            for table in doc.tables:
                table_md = self._table_to_markdown(table)
                if table_md:
                    content_parts.append(table_md)
        
        except Exception as e:
            logger.error(f"DOCX conversion failed: {filename}, error: {e}")
            raise ValueError(f"DOCX 转换失败: {str(e)}")
        
        content = "\n\n".join(content_parts)
        
        if not content.strip():
            warnings.append("DOCX 未提取到任何内容")
        
        if ocr_count > 0:
            warnings.append(f"使用 OCR 识别了 {ocr_count} 张图片")
        
        return ConversionResult(
            content=content, 
            warnings=warnings,
            ocr_used=ocr_count > 0,
            ocr_pages=ocr_count
        )
    
    def _ocr_paragraph_images(
        self, 
        para, 
        warnings: List[str],
        ocr_lang: Optional[str] = None
    ) -> str:
        """对段落中的图片进行 OCR"""
        try:
            from file_converter.ocr.engine import get_ocr_engine
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            
            ocr_engine = get_ocr_engine()
            if not ocr_engine.is_available():
                return ""
            
            # 尝试获取图片数据
            for run in para.runs:
                if run._element.xpath('.//w:drawing'):
                    # 获取内联图片
                    for inline in run._element.xpath('.//a:blip/@r:embed', 
                        namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                                    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}):
                        try:
                            # 获取图片二进制
                            image_part = para.part.related_parts.get(inline)
                            if image_part:
                                image_blob = image_part.blob
                                text = ocr_engine.recognize(image_blob, lang=ocr_lang)
                                if text.strip():
                                    return text.strip()
                        except Exception as e:
                            logger.warning(f"Failed to extract image from DOCX: {e}")
            
            return ""
        
        except Exception as e:
            logger.warning(f"Failed to OCR image in DOCX: {e}")
            return ""
    
    def _process_paragraph(self, para) -> str:
        """处理段落"""
        text = para.text.strip()
        if not text:
            return ""
        
        # 检查标题样式
        style_name = para.style.name if para.style else ""
        heading_level = self.HEADING_STYLES.get(style_name, 0)
        
        if heading_level:
            return "#" * heading_level + " " + text
        
        # 检查列表
        if self._is_list_paragraph(para):
            return self._format_list_item(para)
        
        # 普通段落：处理内联格式
        formatted_text = self._format_runs(para)
        return formatted_text
    
    def _is_list_paragraph(self, para) -> bool:
        """检查段落是否是列表项"""
        try:
            numPr = para._element.xpath('.//w:numPr')
            return len(numPr) > 0
        except Exception:
            return False
    
    def _format_list_item(self, para) -> str:
        """格式化列表项"""
        text = self._format_runs(para)
        
        # 尝试获取列表级别
        try:
            ilvl = para._element.xpath('.//w:ilvl/@w:val')
            level = int(ilvl[0]) if ilvl else 0
        except Exception:
            level = 0
        
        indent = "  " * level
        
        # 检查是否是有序列表
        try:
            numId = para._element.xpath('.//w:numId/@w:val')
            if numId:
                return f"{indent}1. {text}"
        except Exception:
            pass
        
        return f"{indent}- {text}"
    
    def _format_runs(self, para) -> str:
        """处理段落中的 runs，转换加粗和斜体"""
        parts: List[str] = []
        
        for run in para.runs:
            text = run.text
            if not text:
                continue
            
            # 处理加粗
            if run.bold:
                text = f"**{text}**"
            
            # 处理斜体
            if run.italic:
                text = f"*{text}*"
            
            # 处理删除线
            if run.font.strike:
                text = f"~~{text}~~"
            
            parts.append(text)
        
        return "".join(parts) if parts else para.text.strip()
    
    def _table_to_markdown(self, table) -> str:
        """将 DOCX 表格转换为 Markdown"""
        if not table.rows:
            return ""
        
        lines: List[str] = []
        col_count = len(table.columns)
        
        for row_idx, row in enumerate(table.rows):
            cells: List[str] = []
            for cell in row.cells:
                # 提取单元格文本
                text = cell.text.strip()
                # 转义管道符和换行
                text = text.replace("|", "\\|").replace("\n", " ")
                cells.append(text)
            
            # 确保列数一致
            while len(cells) < col_count:
                cells.append("")
            
            lines.append("| " + " | ".join(cells[:col_count]) + " |")
            
            # 在第一行后添加分隔行
            if row_idx == 0:
                lines.append("| " + " | ".join(["---"] * col_count) + " |")
        
        return "\n".join(lines)
